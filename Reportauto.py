import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from io import BytesIO
import openpyxl
import os

st.set_page_config(page_title="Phân tích Điện lực", layout="wide")
st.title("🔍 Ứng dụng kiểm tra hệ thống đo đếm điện")

# Tabs
tabs = st.tabs(["🛠️ Xử lý dữ liệu", "📊 Đánh giá và phân tích", "📈 So sánh dữ liệu"])

# ================= Tab 1: Xử lý dữ liệu =================
with tabs[0]:
    st.header("🛠️ Xử lý dữ liệu kết quả kiểm tra")

    uploaded_source = st.file_uploader("📂 Tải lên file nguồn (Kết quả kiểm tra)", type=["xlsx"], key="source")
    uploaded_template = st.file_uploader("📂 Tải lên file mẫu tổng hợp (Tong hop ket qua ktra dinh ky)", type=["xlsx"], key="template")

    if uploaded_source and uploaded_template:
        from tempfile import NamedTemporaryFile
        with NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(uploaded_source.read())
            tmp_path = tmp.name
        wb_source = openpyxl.load_workbook(tmp_path, data_only=True)
        ws_source = wb_source.active

        wb_target = openpyxl.load_workbook(uploaded_template)
        ws_target = wb_target["KQUA"]

        for row in range(5, 77):
            for col in range(3, 9):
                value = ws_source.cell(row=row, column=col).value
                ws_target.cell(row=row, column=col).value = value

        today_str = datetime.now().strftime("%d%m%Y")
        output_folder = r"C:\Project_Tiennm\Bao cao ktra dinh ky"
        os.makedirs(output_folder, exist_ok=True)
        output_path = os.path.join(output_folder, f"Tong_hop_ket_qua_den_ngay_{today_str}.xlsx")

        wb_target.save(output_path)

        st.success(f"✅ Đã xử lý và lưu file tại: {output_path}")

        # Cho phép tải trực tiếp file sau xử lý
        output_buffer = BytesIO()
        wb_target.save(output_buffer)
        output_buffer.seek(0)

        st.download_button(
            label="📥 Tải file tổng hợp đã xử lý",
            data=output_buffer,
            file_name=f"Tong_hop_ket_qua_den_ngay_{today_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# ================= Tab 2: Đánh giá và phân tích =================
with tabs[1]:
    st.header("📊 Đánh giá và phân tích kết quả kiểm tra")

    uploaded_result = st.file_uploader("📂 Tải lên file tổng hợp kết quả đã xử lý", type=["xlsx"], key="result")

    if uploaded_result:
        df = pd.read_excel(uploaded_result, sheet_name="Tong hop luy ke", header=None)
        df_cleaned = df.iloc[4:].copy()
        df_cleaned.columns = [
            "STT", "Điện lực", "1P_GT", "1P_TT", "3P_GT", "3P_TT",
            "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"
        ]
        df_cleaned = df_cleaned[df_cleaned["Điện lực"].notna()]
        cols_to_numeric = ["1P_GT", "1P_TT", "3P_GT", "3P_TT", "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]
        df_cleaned[cols_to_numeric] = df_cleaned[cols_to_numeric].apply(pd.to_numeric, errors='coerce')
        df_cleaned["Tỷ lệ"] = df_cleaned["Tỷ lệ"] * 100

        total_current = df_cleaned["Tổng công tơ"].sum()
        total_plan = df_cleaned["Kế hoạch"].sum()
        current_date = datetime.now()
        days_passed = (current_date - datetime(2025, 1, 1)).days
        days_total = (datetime(2025, 9, 30) - datetime(2025, 1, 1)).days
        avg_per_day = total_current / days_passed
        forecast_total = avg_per_day * days_total
        forecast_ratio = forecast_total / total_plan

        st.metric("Tổng đã thực hiện", f"{total_current:,}")
        st.metric("Kế hoạch", f"{total_plan:,}")
        st.metric("Tốc độ TB/ngày", f"{avg_per_day:.2f}")
        st.metric("Dự báo đến 30/09/2025", f"{int(forecast_total):,}")
        st.metric("Tỷ lệ dự báo", f"{forecast_ratio*100:.2f}%")

        # Phân tích Top 3 / Bottom 3
        df_sorted = df_cleaned.sort_values(by="Tỷ lệ", ascending=False)
        top_3 = df_sorted.head(3)
        bottom_3 = df_sorted.tail(3)

        # Biểu đồ tổng hợp
        fig, ax = plt.subplots(figsize=(10, 5))
        bars = ax.bar(df_cleaned["Điện lực"], df_cleaned["Tỷ lệ"])
        ax.set_ylabel("Tỷ lệ hoàn thành (%)")
        ax.set_title("Tỷ lệ hoàn thành kế hoạch theo Điện lực")
        ax.tick_params(axis='x', rotation=45)
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
        st.pyplot(fig)

        # Biểu đồ Top 3
        fig_top, ax_top = plt.subplots(figsize=(6, 4))
        bars_top = ax_top.bar(top_3["Điện lực"], top_3["Tỷ lệ"])
        ax_top.set_title("Top 3 Điện lực tỷ lệ cao nhất")
        for bar in bars_top:
            height = bar.get_height()
            ax_top.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                            xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
        st.pyplot(fig_top)

        # Biểu đồ Bottom 3
        fig_bot, ax_bot = plt.subplots(figsize=(6, 4))
        bars_bot = ax_bot.bar(bottom_3["Điện lực"], bottom_3["Tỷ lệ"])
        ax_bot.set_title("Bottom 3 Điện lực tỷ lệ thấp nhất")
        for bar in bars_bot:
            height = bar.get_height()
            ax_bot.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                            xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
        st.pyplot(fig_bot)

        # Đánh giá nâng cao
        st.subheader("📌 Đánh giá tổng thể")
        if forecast_ratio * 100 >= 100:
            st.success("✅ Dự báo sẽ ĐẠT kế hoạch tổng thể đến 30/09/2025.")
        else:
            st.error("❌ Dự báo sẽ KHÔNG ĐẠT kế hoạch tổng thể đến 30/09/2025.")

        # Đánh giá từng Điện lực đạt / không đạt kế hoạch
        st.subheader("📋 Đánh giá theo từng Điện lực")
        df_cleaned["Dự báo tỷ lệ"] = df_cleaned["Tỷ lệ"] * (days_total / days_passed)
        df_cleaned["Kết luận"] = df_cleaned["Dự báo tỷ lệ"].apply(lambda x: "Đạt" if x >= 100 else "Không đạt")
        st.dataframe(df_cleaned[["Điện lực", "Tổng công tơ", "Kế hoạch", "Tỷ lệ", "Dự báo tỷ lệ", "Kết luận"]])

        # Xuất Word có biểu đồ
        def generate_analysis_docx():
            doc = Document()
            doc.add_heading('BÁO CÁO PHÂN TÍCH KẾT QUẢ KIỂM TRA', 0)
            doc.add_paragraph(f"Tổng công tơ đã thực hiện: {total_current:,}")
            doc.add_paragraph(f"Kế hoạch tổng: {total_plan:,}")
            doc.add_paragraph(f"Tốc độ trung bình/ngày: {avg_per_day:.2f} công tơ/ngày")
            doc.add_paragraph(f"Dự báo đến 30/09/2025: {int(forecast_total):,} công tơ")
            doc.add_paragraph(f"Tỷ lệ dự báo: {forecast_ratio*100:.2f}%")
            if forecast_ratio * 100 >= 100:
                doc.add_paragraph("✅ Dự báo sẽ ĐẠT kế hoạch tổng thể đến 30/09/2025.")
            else:
                doc.add_paragraph("❌ Dự báo sẽ KHÔNG ĐẠT kế hoạch tổng thể đến 30/09/2025.")

            # Đánh giá từng Điện lực
            doc.add_heading("I. KẾT QUẢ THỰC HIỆN TỪNG ĐIỆN LỰC", level=1)
            table_eval = doc.add_table(rows=1, cols=6)
            table_eval.style = 'Table Grid'
            hdr = table_eval.rows[0].cells
            hdr[0].text = "Điện lực"
            hdr[1].text = "Tổng công tơ"
            hdr[2].text = "Kế hoạch"
            hdr[3].text = "Tỷ lệ hiện tại"
            hdr[4].text = "Dự báo tỷ lệ"
            hdr[5].text = "Kết luận"
            for _, row in df_cleaned.iterrows():
                r = table_eval.add_row().cells
                r[0].text = str(row['Điện lực'])
                r[1].text = f"{int(row['Tổng công tơ']) if pd.notna(row['Tổng công tơ']) and not pd.isna(row['Tổng công tơ']) else 'N/A'}"
                r[2].text = f"{int(row['Kế hoạch']) if pd.notna(row['Kế hoạch']) and not pd.isna(row['Kế hoạch']) else 'N/A'}"
                r[3].text = f"{row['Tỷ lệ']:.2f}%" if pd.notna(row['Tỷ lệ']) and not pd.isna(row['Tỷ lệ']) else 'N/A'
                r[4].text = f"{row['Dự báo tỷ lệ']:.2f}%" if pd.notna(row['Dự báo tỷ lệ']) and not pd.isna(row['Dự báo tỷ lệ']) else 'N/A'
                r[5].text = row['Kết luận'] if pd.notna(row['Kết luận']) else 'N/A'

                if row['Kết luận'] == 'Không đạt':
                    for cell in r:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
            

            img_buf_all, img_buf_top, img_buf_bot = BytesIO(), BytesIO(), BytesIO()
            fig.savefig(img_buf_all, format="png")
            fig_top.savefig(img_buf_top, format="png")
            fig_bot.savefig(img_buf_bot, format="png")
            img_buf_all.seek(0)
            img_buf_top.seek(0)
            img_buf_bot.seek(0)

            doc.add_heading("II. BIỂU ĐỒ TỔNG HỢP", level=1)
            doc.add_picture(img_buf_all, width=Inches(6))
            doc.add_heading("III. TOP 3", level=1)
            doc.add_picture(img_buf_top, width=Inches(5))
            table_top = doc.add_table(rows=1, cols=4)
            table_top.style = 'Table Grid'
            hdr_cells = table_top.rows[0].cells
            hdr_cells[0].text = "Điện lực"
            hdr_cells[1].text = "Tổng công tơ"
            hdr_cells[2].text = "Kế hoạch"
            hdr_cells[3].text = "Tỷ lệ"
            for _, row in top_3.iterrows():
                row_cells = table_top.add_row().cells
                row_cells[0].text = str(row['Điện lực'])
                row_cells[1].text = f"{int(row['Tổng công tơ']):,}"
                row_cells[2].text = f"{int(row['Kế hoạch']):,}"
                row_cells[3].text = f"{row['Tỷ lệ']:.2f}%"
            doc.add_heading("IV. BOTTOM 3", level=1)
            doc.add_picture(img_buf_bot, width=Inches(5))
            table_bot = doc.add_table(rows=1, cols=4)
            table_bot.style = 'Table Grid'
            hdr_cells = table_bot.rows[0].cells
            hdr_cells[0].text = "Điện lực"
            hdr_cells[1].text = "Tổng công tơ"
            hdr_cells[2].text = "Kế hoạch"
            hdr_cells[3].text = "Tỷ lệ"
            for _, row in bottom_3.iterrows():
                row_cells = table_bot.add_row().cells
                row_cells[0].text = str(row['Điện lực'])
                row_cells[1].text = f"{int(row['Tổng công tơ']):,}"
                row_cells[2].text = f"{int(row['Kế hoạch']):,}"
                row_cells[3].text = f"{row['Tỷ lệ']:.2f}%"

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        analysis_file = generate_analysis_docx()
        st.download_button("📄 Tải báo cáo Word phân tích", data=analysis_file, file_name="Bao_cao_Phan_tich_DienLuc.docx")

# ================= Tab 3: So sánh dữ liệu =================
with tabs[2]:
    st.header("📈 So sánh dữ liệu hiện tại và quá khứ")

    uploaded_old = st.file_uploader("📂 Tải lên file dữ liệu cũ", type=["xlsx"], key="old")
    uploaded_new = st.file_uploader("📂 Tải lên file dữ liệu hiện tại", type=["xlsx"], key="new")

    if uploaded_old and uploaded_new:
        df_old = pd.read_excel(uploaded_old, sheet_name="Tong hop luy ke", header=None)
        df_new = pd.read_excel(uploaded_new, sheet_name="Tong hop luy ke", header=None)

        def clean_dataframe(df):
            df_cleaned = df.iloc[4:].copy()
            df_cleaned.columns = [
                "STT", "Điện lực", "1P_GT", "1P_TT", "3P_GT", "3P_TT",
                "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"
            ]
            df_cleaned = df_cleaned[df_cleaned["Điện lực"].notna()]
            cols_to_numeric = ["1P_GT", "1P_TT", "3P_GT", "3P_TT", "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]
            df_cleaned[cols_to_numeric] = df_cleaned[cols_to_numeric].apply(pd.to_numeric, errors='coerce')
            df_cleaned["Tỷ lệ"] = df_cleaned["Tỷ lệ"] * 100
            return df_cleaned

        df_old_clean = clean_dataframe(df_old)
        df_new_clean = clean_dataframe(df_new)

        df_compare = pd.merge(df_old_clean, df_new_clean, on="Điện lực", suffixes=("_Cũ", "_Mới"))
        df_compare["Chênh lệch Tổng công tơ"] = df_compare["Tổng công tơ_Mới"] - df_compare["Tổng công tơ_Cũ"]
        df_compare["Chênh lệch Tỷ lệ"] = df_compare["Tỷ lệ_Mới"] - df_compare["Tỷ lệ_Cũ"]

        st.subheader("📋 Bảng so sánh tổng hợp")
        st.dataframe(df_compare)

        st.subheader("📈 Biểu đồ so sánh Tỷ lệ hoàn thành")
        fig, ax = plt.subplots(figsize=(10,5))
        x = range(len(df_compare))
        ax.bar(x, df_compare["Tỷ lệ_Cũ"], width=0.4, label="Tỷ lệ Cũ", align="center")
        ax.bar([i + 0.4 for i in x], df_compare["Tỷ lệ_Mới"], width=0.4, label="Tỷ lệ Mới", align="center")
        ax.set_xticks([i + 0.2 for i in x])
        ax.set_xticklabels(df_compare["Điện lực"], rotation=45, ha='right')
        ax.set_ylabel("Tỷ lệ hoàn thành (%)")
        ax.legend()
        st.pyplot(fig)

        chart_buffer = BytesIO()
        fig.savefig(chart_buffer, format="png")
        chart_buffer.seek(0)

        def generate_compare_docx():
            doc = Document()
            doc.add_heading("BÁO CÁO SO SÁNH KẾT QUẢ KIỂM TRA", 0)
            doc.add_paragraph(f"Ngày so sánh: {datetime.now().strftime('%d/%m/%Y')}")

            doc.add_heading("I. BẢNG SO SÁNH", level=1)
            table = doc.add_table(rows=1, cols=7)
            hdr_cells = table.rows[0].cells
            headers = ["Điện lực", "Tổng công tơ Cũ", "Tổng công tơ Mới", "Chênh lệch Công tơ", "Tỷ lệ Cũ", "Tỷ lệ Mới", "Chênh lệch Tỷ lệ"]
            for i, text in enumerate(headers):
                hdr_cells[i].text = text

            for _, row in df_compare.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row['Điện lực'])
                cells[1].text = f"{row['Tổng công tơ_Cũ']:,}"
                cells[2].text = f"{row['Tổng công tơ_Mới']:,}"
                cells[3].text = f"{row['Chênh lệch Tổng công tơ']:,}"
                cells[4].text = f"{row['Tỷ lệ_Cũ']:.2f}%"
                cells[5].text = f"{row['Tỷ lệ_Mới']:.2f}%"
                cells[6].text = f"{row['Chênh lệch Tỷ lệ']:.2f}%"

                if row['Chênh lệch Tỷ lệ'] > 0:
                    for cell in cells:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif row['Chênh lệch Tỷ lệ'] < 0:
                    for cell in cells:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

            doc.add_heading("II. BIỂU ĐỒ SO SÁNH", level=1)
            doc.add_picture(chart_buffer, width=Inches(6))

            doc.add_heading("III. NHẬN XÉT", level=1)
            for _, row in df_compare.iterrows():
                if row['Chênh lệch Tỷ lệ'] > 0:
                    doc.add_paragraph(f"- {row['Điện lực']}: Có cải thiện so với kỳ trước.")
                elif row['Chênh lệch Tỷ lệ'] < 0:
                    doc.add_paragraph(f"- {row['Điện lực']}: Có xu hướng giảm hiệu quả.")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        compare_word_file = generate_compare_docx()
        st.download_button("📄 Tải báo cáo Word so sánh", data=compare_word_file, file_name="Bao_cao_So_Sanh_DienLuc.docx")

        def generate_compare_excel():
            output = BytesIO()
            df_compare.to_excel(output, index=False)
            output.seek(0)
            return output

        compare_excel_file = generate_compare_excel()
        st.download_button("📊 Tải file Excel so sánh", data=compare_excel_file, file_name="So_sanh_DienLuc.xlsx")
